"""
Word 報告文件匯出模組（繁體中文版）
含圖表結語分析
"""
import io
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color_hex})
    shading.append(shading_elem)


def format_table(table, header_color='1a5276'):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in table.rows[0].cells:
        set_cell_shading(cell, header_color)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)
                run.font.bold = True
    for i, row in enumerate(table.rows[1:], 1):
        for cell in row.cells:
            if i % 2 == 0:
                set_cell_shading(cell, 'eaf2f8')
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        elem = borders.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): 'single', qn('w:sz'): '4',
            qn('w:space'): '0', qn('w:color'): 'cccccc'})
        borders.append(elem)
    tblPr.append(borders)


def add_chart_conclusion(doc, text):
    """加入圖表結語（帶灰底框）"""
    p = doc.add_paragraph()
    run = p.add_run('【分析結語】')
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1a, 0x52, 0x76)
    run = p.add_run(f' {text}')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x34, 0x49, 0x5e)
    # 加入底色
    pPr = p._element.get_or_add_pPr()
    shading = pPr.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'F0F3F5'})
    pPr.append(shading)
    p.paragraph_format.space_after = Pt(12)


def generate_docx(results, params, chart_images=None, chart_conclusions=None):
    if chart_conclusions is None:
        chart_conclusions = {}

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

    for i in range(1, 4):
        hs = doc.styles[f'Heading {i}']
        hs.font.color.rgb = RGBColor(0x1a, 0x52, 0x76)
        hs.font.name = 'Arial'
        hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

    proj = params.get('project_name', '冷卻水塔分析報告')
    engineer = params.get('engineer', '-')
    date_str = params.get('date', datetime.now().strftime('%Y-%m-%d'))

    # ===== 封面 =====
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('冷卻水塔\n水蒸發量分析報告')
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1a, 0x52, 0x76)
    run.bold = True

    doc.add_paragraph()
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run('_' * 60)
    run.font.color.rgb = RGBColor(0x1a, 0x52, 0x76)

    doc.add_paragraph()
    for label, value in [('專案名稱', proj), ('計算人員', engineer), ('日期', date_str)]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f'{label}：')
        r1.font.color.rgb = RGBColor(0x1a, 0x52, 0x76)
        r1.bold = True
        r2 = p.add_run(value)
        r2.font.size = Pt(12)
    doc.add_page_break()

    # ===== 目錄 =====
    doc.add_heading('目錄', level=1)
    for item in ['1. 輸入參數', '2. 計算方法說明', '3. 計算結果', '4. 水損失分布',
                  '5. 環境空氣性質', '6. 圖表分析', '7. 結論與建議', '附錄 A：公式參考']:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)
    doc.add_page_break()

    # ===== 1. 輸入參數 =====
    doc.add_heading('1. 輸入參數', level=1)
    doc.add_paragraph('以下為本次冷卻水塔蒸發量分析所使用的參數，所有數值均基於設計條件或現場量測數據。')

    tower_label = "逆流式" if params.get('tower_type') == 'counter' else "橫流式"
    season_label = "夏季 (C=1.0)" if params.get('season') == 'summer' else "冬季 (C=0.8)"
    param_rows = [
        ('循環水量 (Q)', f"{params['Q']:.1f}", 'm³/h'),
        ('進水溫度 (T_in)', f"{params['T_in']:.1f}", '°C'),
        ('出水溫度 (T_out)', f"{params['T_out']:.1f}", '°C'),
        ('溫差 (ΔT)', f"{params['T_in'] - params['T_out']:.1f}", '°C'),
        ('乾球溫度', f"{params['Tdb']:.1f}", '°C'),
        ('濕球溫度', f"{params['Twb']:.1f}", '°C'),
        ('空氣流量 (G)', f"{params['G']:.1f}", 'm³/h'),
        ('填料特性係數 (KaV/L)', f"{params['KaV_L']:.2f}", '-'),
        ('濃縮倍數 (COC)', f"{params['COC']:.1f}", '-'),
        ('水塔類型', tower_label, '-'),
        ('季節', season_label, '-'),
    ]
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = '參數'; hdr[1].text = '數值'; hdr[2].text = '單位'
    for name, val, unit in param_rows:
        row = table.add_row().cells
        row[0].text = name; row[1].text = val; row[2].text = unit
    format_table(table)

    # ===== 2. 計算方法 =====
    doc.add_heading('2. 計算方法說明', level=1)
    doc.add_heading('2.1 簡易經驗公式', level=2)
    doc.add_paragraph('E = C × ΔT × Q / 600')
    doc.add_paragraph('其中 E 為蒸發水量 (m³/h)，C 為蒸發係數（夏季 1.0、冬季 0.8），ΔT 為進出水溫差 (°C)，Q 為循環水量 (m³/h)。')
    doc.add_heading('2.2 Merkel 焓差法', level=2)
    doc.add_paragraph('Merkel 法利用飽和空氣焓值與實際空氣焓值之差進行數值積分，考慮了空氣與水之間的質量傳遞過程，在有空氣流量數據時可提供更精確的結果。')
    doc.add_paragraph('關鍵參數包括氣水比 (L/G)、Merkel 數 (KaV/L) 以及逼近溫度。')
    doc.add_heading('2.3 水損失組成', level=2)
    doc.add_paragraph(
        f'總補給水量 = 蒸發損失 + 飛濺損失 + 排污損失。'
        f'飛濺損失率：{results["splash_rate_pct"]:.1f}%（依水塔類型而定）。'
        f'排污損失依濃縮倍數計算：排污量 = 蒸發量 / (COC - 1)。')

    # ===== 3. 計算結果 =====
    doc.add_heading('3. 計算結果', level=1)
    md = results['merkel_details']
    result_rows = [
        ('氣水比 (L/G)', f"{md['LG_ratio']:.3f}", '-'),
        ('Merkel 數', f"{md['merkel_number']:.3f}", '-'),
        ('蒸發量 - 簡易公式', f"{results['E_simple']:.4f}", 'm³/h'),
        ('蒸發量 - Merkel 法', f"{results['E_merkel']:.4f}", 'm³/h'),
        ('飛濺損失', f"{results['E_splash']:.4f}", 'm³/h'),
        ('排污損失', f"{results['E_blowdown']:.4f}", 'm³/h'),
        ('總補給水量', f"{results['E_total']:.4f}", 'm³/h'),
        ('冷卻效率', f"{md['efficiency']:.1f}", '%'),
        ('逼近溫度', f"{md['T_approach']:.1f}", '°C'),
        ('冷卻範圍', f"{md['T_range']:.1f}", '°C'),
    ]
    table2 = doc.add_table(rows=1, cols=3)
    hdr2 = table2.rows[0].cells
    hdr2[0].text = '項目'; hdr2[1].text = '數值'; hdr2[2].text = '單位'
    for name, val, unit in result_rows:
        row = table2.add_row().cells
        row[0].text = name; row[1].text = val; row[2].text = unit
    format_table(table2)

    # ===== 4. 水損失分布 =====
    doc.add_heading('4. 水損失分布', level=1)
    table3 = doc.add_table(rows=1, cols=3)
    hdr3 = table3.rows[0].cells
    hdr3[0].text = '損失類型'; hdr3[1].text = '流量 (m³/h)'; hdr3[2].text = '佔比 (%)'
    for name, flow, pct in [
        ('蒸發損失', results['E_merkel'], results['evap_pct']),
        ('飛濺損失', results['E_splash'], results['splash_pct']),
        ('排污損失', results['E_blowdown'], results['blowdown_pct']),
        ('合計', results['E_total'], 100.0)]:
        row = table3.add_row().cells
        row[0].text = name; row[1].text = f"{flow:.4f}"; row[2].text = f"{pct:.1f}"
    format_table(table3)

    # ===== 5. 空氣性質 =====
    doc.add_heading('5. 環境空氣性質', level=1)
    table4 = doc.add_table(rows=1, cols=4)
    hdr4 = table4.rows[0].cells
    hdr4[0].text = '性質'; hdr4[1].text = '入口'; hdr4[2].text = '出口'; hdr4[3].text = '單位'
    for r in [
        ('溫度', f"{params['Tdb']:.1f}", f"{md['T_air_out']:.1f}", '°C'),
        ('比濕度', f"{results['W_inlet']:.5f}", f"{md['W_air_out']:.5f}", 'kg/kg'),
        ('焓值', f"{results['h_inlet']:.2f}", f"{md['h_air_out']:.2f}", 'kJ/kg'),
        ('相對濕度', f"{results['RH']:.1f}", '-', '%')]:
        row = table4.add_row().cells
        for i, v in enumerate(r):
            row[i].text = v
    format_table(table4)

    # ===== 6. 圖表分析（含結語）=====
    if chart_images:
        doc.add_heading('6. 圖表分析', level=1)

        chart_info = [
            ('water_loss_pie', '圖 6-1：水損失分布比例'),
            ('cooling_curve', '圖 6-2：冷卻曲線（焓值 vs 溫度）'),
            ('coc_trend', '圖 6-3：濃縮倍數 vs 補給水量'),
            ('temp_enthalpy', '圖 6-4：溫度 - 焓值關係圖'),
        ]

        for key, title in chart_info:
            if key in chart_images and chart_images[key]:
                try:
                    doc.add_heading(title, level=2)
                    img_stream = io.BytesIO(chart_images[key])
                    doc.add_picture(img_stream, width=Inches(5.5))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # 圖表結語
                    if key in chart_conclusions:
                        add_chart_conclusion(doc, chart_conclusions[key])
                except Exception:
                    pass

    # ===== 7. 結論與建議 =====
    doc.add_heading('7. 結論與建議', level=1)
    daily = results['E_total'] * 24
    monthly = daily * 30
    doc.add_paragraph(f'總補給水量需求為 {results["E_total"]:.4f} m³/h，約相當於每日 {daily:.1f} m³ 或每月 {monthly:.0f} m³。')
    if md['efficiency'] > 70:
        doc.add_paragraph('冷卻效率良好，超過 70%，水塔運行狀態佳。')
    elif md['efficiency'] > 50:
        doc.add_paragraph('冷卻效率中等，建議優化風量或檢查填料狀況。')
    else:
        doc.add_paragraph('冷卻效率低於 50%，請檢查填料狀態、風扇性能及擋水器。')
    if params['COC'] < 3:
        doc.add_paragraph(f'目前濃縮倍數為 {params["COC"]:.1f}，提高至 4-6 可顯著減少排污用水量。')

    # ===== 附錄 =====
    doc.add_page_break()
    doc.add_heading('附錄 A：公式參考', level=1)
    doc.add_heading('A.1 飽和水蒸氣壓力', level=2)
    doc.add_paragraph('Ps = 0.61078 × exp(17.27 × T / (T + 237.3))  [kPa]')
    doc.add_heading('A.2 比濕度', level=2)
    doc.add_paragraph('W = 0.622 × Pw / (P - Pw)  [kg 水 / kg 乾空氣]')
    doc.add_heading('A.3 濕空氣焓值', level=2)
    doc.add_paragraph('h = 1.006 × T + W × (2501 + 1.86 × T)  [kJ/kg 乾空氣]')
    doc.add_heading('A.4 冷卻塔效率', level=2)
    doc.add_paragraph('效率 = 冷卻範圍 / (冷卻範圍 + 逼近溫度) × 100%')
    doc.add_paragraph('冷卻範圍 = T_in - T_out；逼近溫度 = T_out - T_wb')
    doc.add_heading('A.5 基本假設', level=2)
    for a in ['大氣壓力：101.325 kPa（海平面）', '水的密度：1000 kg/m³',
              '水的比熱：4.186 kJ/(kg·°C)', '空氣密度：1.2 kg/m³（近似值）', '假設穩態運行條件']:
        doc.add_paragraph(a, style='List Bullet')

    for section in doc.sections:
        hp = section.header.paragraphs[0]
        hp.text = proj
        hp.style.font.size = Pt(8)
        hp.style.font.color.rgb = RGBColor(0x1a, 0x52, 0x76)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
